# process_uploaded.py
import os
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

UPLOAD_DIR = Path('uploads')
OUT_DIR = Path('processed')
OUT_DIR.mkdir(exist_ok=True)

# Basit ayarlar: sabit bir format uyguluyoruz. İleri konfigürasyon için bu dosyayı genişletebilirsiniz.
DEFAULT_FONT = 'Times New Roman'
DEFAULT_SIZE = 12
MARGINS_INCH = (1,1,1,1)  # top, bottom, left, right

def apply_font_to_paragraph(p, font_name=None, size_pt=None, bold=None, italic=None):
    for r in p.runs:
        if font_name:
            r.font.name = font_name
        if size_pt:
            r.font.size = Pt(size_pt)
        if bold is not None:
            r.font.bold = bold
        if italic is not None:
            r.font.italic = italic

def process_file(path: Path):
    print('Processing', path)
    doc = Document(path)

    # margins
    top, bottom, left, right = MARGINS_INCH
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

    # default font/size
    for p in doc.paragraphs:
        apply_font_to_paragraph(p, font_name=DEFAULT_FONT, size_pt=DEFAULT_SIZE)

    # headings: basit örnek - Heading 1 için
    for p in doc.paragraphs:
        try:
            if p.style.name and 'Heading 1' in p.style.name:
                apply_font_to_paragraph(p, font_name='Times New Roman', size_pt=16, bold=True)
        except Exception:
            pass

    outp = OUT_DIR / (path.stem + '_formatted.docx')
    doc.save(outp)
    print('Saved', outp)

def main():
    if not UPLOAD_DIR.exists():
        print('No uploads folder found. Nothing to do.')
        return
    files = list(UPLOAD_DIR.glob('*.docx'))
    if not files:
        print('No .docx files in uploads/. Nothing to do.')
        return
    for f in files:
        try:
            process_file(f)
        except Exception as e:
            print('Error processing', f, e)

if __name__ == '__main__':
    main()
